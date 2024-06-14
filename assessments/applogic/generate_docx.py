import json
import os
from tempfile import NamedTemporaryFile

from docx import Document
from docx.enum.style import WD_BUILTIN_STYLE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from django.core.files import File


def get_assessment_models():  # pragma: no cover

    from .. import models as assessment_models

    return assessment_models


def add_checkbox(checked=False):
    return "☑" if checked else "☐"


def set_table_row_height(table, height=10):  # pragma: no cover

    for row in table.rows:
        row.height = Pt(height)


def define_criteria_row(table, index, criteria, form_dict, key):  # pragma: no cover

    row = table.add_row()
    row.cells[0].text = str(index)
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[0].width = Pt(30)
    row.cells[1].width = Pt(120)
    row.cells[2].width = Pt(60)
    paragraph = row.cells[1].add_paragraph()
    for string in criteria:
        paragraph.add_run(string)
    paragraph.paragraph_format.left_index = Inches(0.25)

    paragraph = row.cells[2].add_paragraph()

    paragraph.add_run(add_checkbox(form_dict.get(key)))
    paragraph.add_run("Yes\n")
    paragraph.add_run(add_checkbox(not form_dict.get(key)))
    paragraph.add_run("No\n")

    if comment := form_dict.get(f"{key}_comment"):
        run = paragraph.add_run("\nComment:\n")
        run.font.bold = True
        paragraph.add_run(comment)

    return row


def define_criteria_row_header(table, title):  # pragma: no cover

    row = table.add_row()
    merged = row.cells[0].merge(row.cells[1])
    merged = merged.merge(row.cells[2])
    merged.text = title
    row.height = Pt(40)
    merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged.paragraphs[0].runs[0].font.bold = True
    merged.alignment = WD_TABLE_ALIGNMENT.CENTER
    return row


def insert_table_rows(table, form, data):  # pragma: no cover
    for line in data:
        if line.get("is_header"):
            define_criteria_row_header(table, line["text"])
        else:
            define_criteria_row(
                table, line["index"], line["criteria"], form, line["key"]
            )


def insert_classification_row(
    table, classification, title, paragraphs, key
):  # pragma: no cover
    row = table.add_row()
    row.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row.cells[0].width = Pt(30)
    row.cells[1].width = Pt(300)
    row.cells[2].width = Pt(30)

    row.cells[0].text = title
    cell = row.cells[1]
    for paragraph in paragraphs:
        cell.add_paragraph(paragraph)

    row.cells[2].text = add_checkbox(classification.get(key))
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return row


def generate_docx_output(form):

    doc = Document()
    data = {}

    # HEADER
    heading_table = doc.add_table(rows=2, cols=1, style="Table Grid")
    heading_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    heading_table.rows[0].cells[0].text = form.name
    heading_table.rows[1].cells[0].text = "Title: Compliance Criticality Assessment"

    heading_table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_table.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    heading_table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True

    heading_table.base_style = WD_BUILTIN_STYLE.TABLE_LIGHT_GRID

    heading = doc.add_heading("1   PURPOSE", level=2)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()
    paragraph = doc.add_paragraph(
        f"""1.1	The purpose of this Compliance Criticality Assessment (CCA) is to provide a documented evaluation that establishes and defines the overall classification (i.e., GxP, SOX, non-GxP, Privacy, etc.) and business impact of {form.name} based on the solution functionality and intended use."""
    )

    doc.add_paragraph()

    paragraph = doc.add_paragraph(
        """1.2 This document is governed by the IT Risk Management SOP."""
    )

    heading = doc.add_heading("2   SOLUTION IDENTIFICATION", level=2)
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()

    table = doc.add_table(rows=7, cols=2, style="Table Grid")
    set_table_row_height(table, 15)

    table.rows[0].cells[0].text = "Solution Name"
    if form.solution_name:
        table.rows[0].cells[1].text = form.solution_name
        table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

    table.rows[1].cells[0].text = "Software Release / Version, as applicable"
    if form.solution_version:
        table.rows[1].cells[1].text = form.solution_version
        table.rows[1].cells[1].paragraphs[0].runs[0].font.bold = True

    table.rows[2].cells[0].text = "Vendor Name"
    if form.vendor_name:
        table.rows[2].cells[1].text = form.vendor_name
        table.rows[2].cells[1].paragraphs[0].runs[0].font.bold = True

    table.rows[3].cells[0].text = "Type of solution"

    cell = table.rows[3].cells[1]
    sub_table = cell.add_table(rows=2, cols=2)

    cell = sub_table.rows[0].cells[0]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.solution_type == "application"))
    paragraph.add_run(" Application")

    cell = sub_table.rows[0].cells[1]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.solution_type == "middleware"))
    paragraph.add_run(" Middleware")

    cell = sub_table.rows[1].cells[0]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.solution_type == "infrastructure"))
    paragraph.add_run(" Infrastructure Platform")

    cell = sub_table.rows[1].cells[1]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.solution_type == "other"))
    paragraph.add_run(" Others")
    if form.other_solution_type:
        paragraph.add_run(f" {form.other_solution_type}")

    table.rows[4].cells[0].text = "Hosting and Type"
    sub_table = table.rows[4].cells[1].add_table(rows=3, cols=2)

    cell = sub_table.rows[0].cells[0]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.hosting_and_type == "third_party"))
    paragraph.add_run(" Third Party")

    cell = sub_table.rows[0].cells[1]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.hosting_and_type == "saas"))
    paragraph.add_run(" Software as a Service(SAAS)")

    cell = sub_table.rows[1].cells[0]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.hosting_and_type == "on_premises"))
    paragraph.add_run(" On-Premises")

    cell = sub_table.rows[1].cells[1]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.hosting_and_type == "paas"))
    paragraph.add_run(" Platform as a Service(PAAS)")

    cell = sub_table.rows[2].cells[0]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.hosting_and_type == "website"))
    paragraph.add_run(" Website (Public)")

    cell = sub_table.rows[2].cells[1]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.hosting_and_type == "iaas"))
    paragraph.add_run(" Infrastructure as a Service(IAAS)")

    table.rows[5].cells[0].text = "Server Host / Location"
    if form.server_host:
        table.rows[5].cells[1].text = form.server_host
        table.rows[5].cells[1].paragraphs[0].runs[0].font.bold = True

    table.rows[6].cells[0].text = "Solution Classification"
    sub_table = table.rows[6].cells[1].add_table(rows=3, cols=1)

    cell = sub_table.rows[0].cells[0]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.solution_classification == "custom"))
    paragraph.add_run(" Custom")

    cell = sub_table.rows[1].cells[0]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.solution_classification == "configurable"))
    paragraph.add_run(" Configurable")

    cell = sub_table.rows[2].cells[0]
    paragraph = cell.paragraphs[0]
    paragraph.add_run(add_checkbox(form.solution_classification == "non_configurable"))
    paragraph.add_run(" Non-configurable (out-of-the-box)")

    for cell in table.columns[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    heading = doc.add_heading("3   SOLUTION DESCRIPTION AND INTENDED USE", level=2)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()

    paragraph = doc.add_paragraph(f"""    3.1   {form.solution_description or "N/A"}""")

    heading = doc.add_heading("4   SOLUTION REGULATORY IMPACT")
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()

    paragraph = doc.add_paragraph("   4.1  GxP Impact")

    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_1 = table.cell(0, 0)
    col_2 = table.cell(0, 1)
    merged_cell = col_1.merge(col_2)
    merged_cell.text = "Criteria"
    table.rows[0].cells[2].text = "Yes / No"

    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

    row = table.add_row()
    merged = row.cells[0].merge(row.cells[1])
    merged = merged.merge(row.cells[2])
    merged.text = "Good Manufacturing Practices (GMP) Impact"
    row.height = Pt(40)
    merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged.paragraphs[0].runs[0].font.bold = True
    merged.alignment = WD_TABLE_ALIGNMENT.CENTER

    current_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(current_dir, "table_content.json")

    with open(file_path, "r") as json_file:
        data = json.load(json_file)

    insert_table_rows(table, form.gxp_impact, data["gxp_impact"])

    doc.add_paragraph(
        "\n4.2	GxP Electronic Records (ER) and Electronic Signatures (ES) Applicability\n"
    )

    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_1 = table.cell(0, 0)
    col_2 = table.cell(0, 1)
    merged_cell = col_1.merge(col_2)
    merged_cell.text = "Criteria"
    table.rows[0].cells[2].text = "Yes / No"

    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

    # START ROW
    insert_table_rows(table, form.gxp_eres, data["gxp_eres"])

    doc.add_paragraph("\n4.3	SOX Impact\n")

    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_1 = table.cell(0, 0)
    col_2 = table.cell(0, 1)
    merged_cell = col_1.merge(col_2)
    merged_cell.text = "Criteria"
    table.rows[0].cells[2].text = "Yes / No"

    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

    insert_table_rows(table, form.gxp_eres, data["sox_impact"])

    doc.add_paragraph("\n4.4	Privacy Impact\n")

    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_1 = table.cell(0, 0)
    col_2 = table.cell(0, 1)
    merged_cell = col_1.merge(col_2)
    merged_cell.text = "Criteria"
    table.rows[0].cells[2].text = "Yes / No"

    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

    # START ROW
    index = "1"
    criteria = """Does the Solution collect, process, and/or disclose Personal Information (or does not prevent the ability to), directly or through a third party, as defined in the Processing EU Personal Data and/ or in the Company standard operating procedure, Processing of European Union Personal Data and Privacy by Design Standard Operating Procedure?"""
    key = "does_solution_collect_person_information"
    row = define_criteria_row(table, index, criteria, form.privacy_impact, key)

    row.cells[1].add_paragraph()
    paragraph = row.cells[1].add_paragraph()
    run = paragraph.add_run(
        """Note: The EU countries are: Austria, Belgium, Bulgari, Croatia, Republic of Cyprus, Czech Republic, Denmark, Estonia, Finland, France, Germany, Greece, Hungary, Ireland, Italy, Latvia, Lithuania, Luxembourg, Malta, Netherlands, Poland, Portugal, Romania, Slovakia, Slovenia, Spain and Sweden."""
    )
    run.italic = True

    # END ROW

    # START ROW
    index = "2"
    criteria = """Will the personal information being collected, processed, or disclosed belong to an individual not residing in an EU country?"""
    key = "will_personal_info_collected_be_from_individual_not_in_eu_country"
    define_criteria_row(table, index, criteria, form.privacy_impact, key)

    # END ROW

    heading = doc.add_heading("5   DATA CLASSIFICATION")
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph("\nNot Examples include data as described in Appendix 2.")

    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    table.rows[0].cells[0].text = "Classification"
    table.rows[0].cells[1].text = "Description"
    table.rows[0].cells[2].text = "Rating"

    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

    for line in data["classifications"]:
        insert_classification_row(
            table,
            form.data_classification,
            line["title"],
            line["paragraphs"],
            line["key"],
        )

    heading = doc.add_heading("6   SOLUTION BUSINESS IMPACT")
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    cell_1, cell_2, cell_3 = table.rows[0].cells
    merged = cell_1.merge(cell_2)
    merged = merged.merge(cell_3)

    paragraph = merged.add_paragraph(
        "The Business Impact is used to determine the following: \n"
    )
    run = paragraph.runs[0]
    run.font.bold = True
    paragraph = merged.add_paragraph("")
    run = paragraph.add_run(
        "\t1. Determine mission/business processes and recovery criticality."
    )
    run.font.bold = True
    run = paragraph.add_run(
        "  Mission/business processes supported by the system are identified and the impact of a system disruption to those processes is determined along with outage impact and estimated downtime.  The downtime should reflect the maximum that an organization can tolerate while still maintaining the mission.\n"
    )
    paragraph = merged.add_paragraph("")

    run = paragraph.add_run("\t2. Identify resource requirements.")
    run.font.bold = True
    run = paragraph.add_run(
        "  Realistic recovery efforts require a thorough evaluation of the resources required to resume mission/business processes and related interdependencies as quickly as possible.  Examples of resources that should be identified include facilities, personnel, equipment, software, data files, system components, and vital records.\n"
    )

    run = paragraph.add_run("\t3. Identify recovery priorities for system resources.")
    run.font.bold = True
    run = paragraph.add_run(
        "  Based upon the results from the previous activities, system resources can more clearly be linked to critical mission/business processes.  Priority levels can be established for sequencing recovery activities and resources.\n"
    )

    row = table.add_row()
    row.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    merged = row.cells[0].merge(row.cells[1])
    merged.text = "Criteria"
    merged.paragraphs[0].runs[0].font.bold = True
    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row.cells[2].text = "Rating"
    row.cells[2].paragraphs[0].runs[0].font.bold = True
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = table.add_row()
    row.cells[0].width = Pt(30)
    row.cells[1].width = Pt(300)
    row.cells[2].width = Pt(60)

    row.cells[0].text = "1"
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = row.cells[1].add_paragraph("")
    run = paragraph.add_run("High ")
    run.font.bold = True
    run = paragraph.add_run("(Patient Safety)")
    run.font.italic = True
    paragraph = row.cells[1].add_paragraph(
        """
        \t - Risk to patent safety (GxP)
        \t - Serious disruption of business with no compensating manual processes available.\n"""
    )
    row.cells[2].text = add_checkbox(form.business_impact.get("business_impact_high"))
    row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = table.add_row()
    row.cells[0].text = "2"
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = row.cells[1].paragraphs[0]
    run = paragraph.add_run("Medium ")
    run.font.bold = True
    run = paragraph.add_run("(Business Mission Critical)")
    run.font.italic = True
    paragraph = row.cells[1].add_paragraph(
        """
        \t(a) revenue impacted
        \t(b) negative customer satisfaction
        \t(c) compliance violation (not Patient safety) and/or
        \t(d) damage to organization's reputation\n"""
    )
    row.cells[2].text = add_checkbox(form.business_impact.get("business_impact_medium"))
    row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = table.add_row()

    row.cells[0].text = "3"
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = row.cells[1].paragraphs[0]
    run = paragraph.add_run("Low ")
    run.font.bold = True
    run = paragraph.add_run("(Business Supporting)")
    run.font.italic = True
    run = paragraph.add_run("Employee productivity degradation")
    row.cells[2].text = add_checkbox(form.business_impact.get("business_impact_low"))
    row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in table.rows[1:]:
        row.cells[0].width = Pt(30)
        row.cells[2].width = Pt(50)

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2, style="Table Grid")
    row = table.rows[0]
    paragraph = row.cells[0].paragraphs[0]
    run = paragraph.add_run("\nComments:\n")
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run.font.bold = True

    row.cells[1].text = form.business_impact.get("business_impact_comment") or "N/A"
    row.cells[0].width = Pt(80)
    row.cells[1].width = Pt(450)
    doc.add_paragraph()

    heading = doc.add_heading("7   SOLUTION CRITICALITY SUMMARY", level=2)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=4, cols=2, style="Table Grid")
    heading = table.rows[0]

    heading.cells[0].text = "Criteria"
    heading.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.cells[0].paragraphs[0].runs[0].font.bold = True

    heading.cells[1].text = "Results"
    heading.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.cells[1].paragraphs[0].runs[0].font.bold = True

    summary = form.summary

    row = table.rows[1]
    row.cells[0].text = "Sec. 4, Solution Regulatory Impact"
    sub_table = row.cells[1].add_table(cols=1, rows=5)

    sub_sub_table = sub_table.rows[0].cells[0].add_table(cols=3, rows=2)

    sub_sub_table.rows[0].cells[
        0
    ].text = f"{add_checkbox(summary.get('summary_regulatory_gxp_impact_gmp'))} GMP"
    sub_sub_table.rows[0].cells[
        1
    ].text = f"{add_checkbox(summary.get('summary_regulatory_gxp_impact_gcp'))} GCP"
    sub_sub_table.rows[0].cells[
        2
    ].text = f"{add_checkbox(summary.get('summary_regulatory_gxp_impact_gxp_indirect'))} GxP Indirect"

    sub_sub_table.rows[1].cells[
        0
    ].text = f"{add_checkbox(summary.get('summary_regulatory_gxp_impact_glp'))} GLP"
    sub_sub_table.rows[1].cells[
        1
    ].text = f"{add_checkbox(summary.get('summary_regulatory_gxp_impact_gvp'))} GVP"
    sub_sub_table.rows[1].cells[
        2
    ].text = (
        f"{add_checkbox(summary.get('summary_regulatory_gxp_impact_non_gxp'))} Non-GxP"
    )

    sub_sub_table = sub_table.rows[0].cells[0].add_table(cols=2, rows=1)
    sub_sub_table.rows[0].cells[
        0
    ].text = f"{add_checkbox(summary.get('summary_regulatory_sox_impact_sox'))} SOX"
    sub_sub_table.rows[0].cells[
        1
    ].text = (
        f"{add_checkbox(summary.get('summary_regulatory_sox_impact_non_sox'))} Non-SOX"
    )

    sub_sub_table = sub_table.rows[0].cells[0].add_table(cols=3, rows=1)
    sub_sub_table.rows[0].cells[
        0
    ].text = f"{add_checkbox(summary.get('summary_regulatory_gxp_eres_er'))} ER"
    sub_sub_table.rows[0].cells[
        1
    ].text = f"{add_checkbox(summary.get('summary_regulatory_gxp_eres_es'))} ES"
    sub_sub_table.rows[0].cells[
        2
    ].text = (
        f"{add_checkbox(summary.get('summary_regulatory_gxp_eres_non_eres'))} Non-ERES"
    )

    sub_sub_table = sub_table.rows[0].cells[0].add_table(cols=2, rows=2)
    sub_sub_table.rows[0].cells[
        0
    ].text = f"{add_checkbox(summary.get('summary_regulatory_privacy_impact_high_privacy'))} High Privacy"
    sub_sub_table.rows[0].cells[
        1
    ].text = f"{add_checkbox(summary.get('summary_regulatory_privacy_impact_medium_privacy'))} Medium Privacy"
    sub_sub_table.rows[1].cells[
        0
    ].text = f"{add_checkbox(summary.get('summary_regulatory_privacy_impact_no_privacy'))} No Privacy"
    sub_sub_table.rows[1].cells[
        1
    ].text = f"{add_checkbox(summary.get('summary_regulatory_privacy_impact_low_privacy'))} Low Privacy"

    sub_sub_table = sub_table.rows[0].cells[0].add_table(cols=1, rows=2)
    sub_sub_table.rows[0].cells[
        0
    ].text = f"{add_checkbox(summary.get('summary_regulatory_impact_administrative_audit_trail_review'))} Administrative Audit Trail Review"
    sub_sub_table.rows[1].cells[
        0
    ].text = f"{add_checkbox(summary.get('summary_regulatory_impact_operational_audit_trail_review'))} Operational Audit Trail Review"

    row = table.rows[2]
    row.cells[0].text = "Sec. 5, Data Classification"

    sub_table = row.cells[1].add_table(cols=2, rows=2)

    sub_table.rows[0].cells[
        0
    ].text = f"{add_checkbox(summary.get('summary_data_classification_secret'))} Secret"
    sub_table.rows[0].cells[
        1
    ].text = f"{add_checkbox(summary.get('summary_data_classification_restricted'))} Restricted"
    sub_table.rows[1].cells[
        0
    ].text = (
        f"{add_checkbox(summary.get('summary_data_classification_internal'))} Internal"
    )
    sub_table.rows[0].cells[
        1
    ].text = f"{add_checkbox(summary.get('summary_data_classification_public'))} Public"

    row = table.rows[3]
    row.cells[0].text = "Sec. 6, Solution Business Impact"

    sub_table = row.cells[1].add_table(rows=1, cols=3)
    sub_table.rows[0].cells[
        0
    ].text = f"{add_checkbox(summary.get('summary_business_impact_high'))} High"
    sub_table.rows[0].cells[
        1
    ].text = f"{add_checkbox(summary.get('summary_business_impact_medium'))} Medium"
    sub_table.rows[0].cells[
        2
    ].text = f"{add_checkbox(summary.get('summary_business_impact_low'))} Low"

    doc.add_paragraph()

    heading = doc.add_heading("8   COMPLIANCE CRITICALITY RATING", level=2)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=1, cols=3, style="Table Grid")

    heading_row = table.rows[0]
    heading_row.cells[0].width = Inches(0.25)
    merged = heading_row.cells[0].merge(heading_row.cells[1])
    merged.text = "Criteria"
    merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged.paragraphs[0].runs[0].font.bold = True

    heading_row.cells[2].text = "Rating"
    heading_row.cells[2].paragraphs[0].runs[0].font.bold = True
    heading_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    rating = form.rating

    row = table.add_row()

    row.cells[0].text = "1"
    row.cells[2].text = add_checkbox(rating.get("rating_significant"))
    paragraph = row.cells[1].paragraphs[0]
    run = paragraph.add_run("Significant\n")
    run.font.bold = True

    run = paragraph.add_run("\t- Regulatory: ")
    run.font.bold = True
    paragraph.add_run("GxP-Direct Impact OR\n")

    run = paragraph.add_run("\t- Data Classification: ")
    run.font.bold = True
    paragraph.add_run("SECRET OR\n")

    row = table.add_row()
    row.cells[0].text = "2"
    row.cells[2].text = add_checkbox(rating.get("rating_moderate"))
    paragraph = row.cells[1].paragraphs[0]
    run = paragraph.add_run("Moderate\n")
    run.font.bold = True

    run = paragraph.add_run("\t- Regulatory: ")
    run.font.bold = True
    paragraph.add_run(
        "GxP-Indirect Impact, SOX, High or Medium Privacy Impact etc. OR\n"
    )

    run = paragraph.add_run("\t- Data Classification: ")
    run.font.bold = True
    paragraph.add_run("INTERNAL OR RESTRICTED\n")

    run = paragraph.add_run("\t- Business Impact: ")
    run.font.bold = True
    paragraph.add_run("High\n")

    row = table.add_row()
    row.cells[0].text = "3"
    row.cells[2].text = add_checkbox(rating.get("rating_minimal"))
    paragraph = row.cells[1].paragraphs[0]
    run = paragraph.add_run("Minimal\n")
    run.font.bold = True

    run = paragraph.add_run("\t- Regulatory: ")
    run.font.bold = True
    paragraph.add_run("Non-GxP Impact, Low Privacy Impact, non-SOX OR \n")

    run = paragraph.add_run("\t- Data Classification: ")
    run.font.bold = True
    paragraph.add_run("PUBLIC OR\n")

    run = paragraph.add_run("\t- Business Impact: ")
    run.font.bold = True
    paragraph.add_run("Medium\n")

    row = table.add_row()
    row.cells[0].text = "4"
    row.cells[2].text = add_checkbox(rating.get("rating_no_compliance_risk"))
    paragraph = row.cells[1].paragraphs[0]
    run = paragraph.add_run("No Compliance Risk\n")
    run.font.bold = True

    run = paragraph.add_run("\t- Regulatory: ")
    run.font.bold = True
    paragraph.add_run("Non-GxP Impact, No Privacy Impact, non-SOX AND \n")

    run = paragraph.add_run("\t- Business Impact: ")
    run.font.bold = True
    paragraph.add_run("Low\n")

    for row in table.rows[1:]:
        row.cells[0].width = Pt(30)
        row.cells[1].width = Pt(400)
        row.cells[2].width = Pt(60)

        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2, style="Table Grid")
    row = table.rows[0]
    paragraph = row.cells[0].paragraphs[0]
    run = paragraph.add_run("\nComments:\n")
    row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run.font.bold = True
    row.cells[0].width = Pt(80)
    row.cells[1].width = Pt(450)

    paragraph = doc.add_paragraph(
        "Note: If Solution has multiple/conflicting CCA Criteria then criteria which indicates worst case scenario (e.g., Scenario 1 above is worse than Scenario 2) should be used."
    )
    paragraph.runs[0].font.italic = True
    doc.add_paragraph()

    heading = doc.add_heading("9   APPENDIX", level=2)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("\t9.1 Appendix 1 - GxP Impact")
    doc.add_paragraph("\t9.2 Appendix 2 - Data Classifications\n")

    heading = doc.add_heading("10   DOCUMENT SIGNATORIES", level=2)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=3, cols=2, style="Table Grid")
    table.rows[0].cells[0].text = "\nSystem Owner\n"
    table.rows[1].cells[0].text = "\nBusiness Owner\n"
    table.rows[2].cells[0].text = "\nIT Risk Management and Compliance\n"

    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].font.italic = True
        row.cells[0].width = Pt(120)
        row.cells[1].width = Pt(300)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading = doc.add_heading("11   DOCUMENT CHANGE HISTORY", level=2)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=2, cols=2, style="Table Grid")
    table.rows[0].cells[0].width = Pt(40)

    table.rows[0].cells[0].text = "VERSION #"
    table.rows[0].cells[1].text = "SUMMARY OF CHANGES (INCLUDING DRIVER/JUSTIFICATION)"

    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
    table.rows[1].cells[0].text = "1.0"
    table.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[1].cells[1].text = "New Document"

    paragraph = doc.add_paragraph("\nAppendix 1 - GxP Impact\n")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.runs[0].font.bold = True
    doc.add_paragraph(
        """Solution is used for Product Quality Control when it is used for the following, among others:
        - laboratory solution
        - raw material testing
        - in-process testing
        - stability
        - environmental monitoring
        """
    )
    doc.add_paragraph(
        """Solution is used to monitor, control, or supervise:
        - Packaging and labeling operations when it is used for the following, among others:
        \t\u2022  Labeling
        \t\u2022  Artwork
        - Production processes and/or materials when it is used for the following, among others:
        \t\u2022  process instrumentation
        \t\u2022  dispensing
        \t\u2022  warehouse
        \t\u2022  inventory
        \t\u2022  inspection
        \t\u2022  serialization
        - GMP facilities/environments or services when it is used for the following, among others:
        \t\u2022  BMS
        \t\u2022  HVAC
        \t\u2022  WFI
        """
    )

    doc.add_paragraph(
        """Example of patient or data subject original documents, data, and records are the following, among others:
        - hospital records
        - clinical charts
        - laboratory notes
        - memoranda
        - subjects' diaries or evaluation checklists
        - pharmacy dispensing records
        - recorded data from automated instruments
        - copies or transcriptions certified after verification as being accurate copies
        - microfiches
        - photographic negatives
        - microfilm or magnetic media
        - x-rays, subject files, and records kept at the pharmacy, at the laboratories and at medico-technical departments
        - computer readable media
        - dictated observations
        - recorded data from automated instruments, or any other data storage medium
        """
    )

    current_dir = os.path.dirname(os.path.abspath(__file__))
    file_path = os.path.join(current_dir, "appendix_2_text.json")

    with open(file_path, "r") as json_file:
        data = json.load(json_file)
        paragraph = doc.add_paragraph("Appendix 2 - Data Classifications")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].font.bold = True
        table = doc.add_table(rows=3, cols=4, style="Table Grid")
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for column in table.columns:
            column.width = Inches(2.5)
        keys = ["secret", "restricted", "internal", "public"]
        for index, key in enumerate(keys):
            cell = table.rows[0].cells[index]
            cell.text = key.upper()
            cell.paragraphs[0].runs[0].font.bold = True
            table.rows[1].cells[index].text = data[key]["text"]
            example_cell = table.rows[2].cells[index]
            run = example_cell.paragraphs[0].add_run("Examples")
            run.font.bold = True
            example_cell.add_paragraph(data[key]["examples"])

    heading = doc.add_heading("CHANGE HISTORY", level=2)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows=4, cols=2, style="Table Grid")
    table.rows[0].cells[0].width = Pt(40)
    table.rows[0].cells[1].width = Pt(800)
    table.rows[0].cells[0].text = "VERSION #"
    table.rows[0].cells[1].text = "SUMMARY OF CHANGES (INCLUDING DRIVER/JUSTIFICATION)"
    for col in table.rows[0].cells:
        col.paragraphs[0].runs[0].font.bold = True

    table.rows[1].cells[0].text = "3.0"
    table.rows[1].cells[1].add_paragraph(
        """- Updated formatting and removed GxP Data Criticality impact classification as it is already covered in impact classification on GxP.
- Updated Section 7 and 8 to be more cohesive.
- Updated Data Classification to align with IT Security data classifications and Bayer language."""
    )

    table.rows[2].cells[0].text = "2.0"
    table.rows[2].cells[1].add_paragraph(
        """Updated form to be aligned with the new Form Template.
Updated Section 4 to separate question pertaining to GCP and GLP Impact.
Updated Section 7 to separate checkbox for GCP and GLP Impact.
Added "No Compliance Risk" as option on the Compliance Criticality Assessment.
Removed Impact column.
Changed section name from Form Execution History to Document Change History
Fixed Appendix section"""
    )

    table.rows[3].cells[0].text = "1.0"
    table.rows[3].cells[1].add_paragraph("New Document to support IT process")

    for row in table.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    with NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx_file:
        doc.save(temp_docx_file.name)

    with open(temp_docx_file.name, "rb") as docx_file:
        assessment_models = get_assessment_models()
        report, _ = assessment_models.Report.objects.get_or_create(cca=form)
        report.output_cca_doc.save("cca_document.docx", File(docx_file))

    os.unlink(temp_docx_file.name)
    return doc
