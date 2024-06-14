import pytest
from docx import Document

from assessments.models import ComplianceCriticalityAssessment, Report
from assessments.tests.factories import ComplianceCriticalityAssessmentFactory

from ..applogic.generate_docx import add_checkbox, generate_docx_output

"""
            TABLE LEGEND BY INDEX
0. Title
1. Solution Identificaiton
2. (4.1) GxP Impact
3. (4.2) GxP Electronic Records (ER) and Electronic Signatures (ES) Applicability
4. (4.3) SOX Impact
5. (4.4) Privacy Impact
6. Data Classification
7. Solution Business Impact
8. Solution Business Impact Comments
9. Solution Criticality Summary
10. Compliance Criticality Rating
11. Compliance Criticality Rating Comments
12. Document Signatories
13. Document Change History
14. Appendix 2
15. Change History

"""


def is_true_checkbox(string):
    return string == "☑"


@pytest.mark.django_db
def test_generate_docx():
    checkbox = {True: "☑", False: "☐"}
    checkbox_choice = {True: "☑Yes", False: "☑No", "": "☑No", None: "☑No"}
    # NOTE: checkbox[True] == "☑", used to check if checkbox is ticked or not

    cca: ComplianceCriticalityAssessment = ComplianceCriticalityAssessmentFactory()

    document = generate_docx_output(cca)

    # Title header
    title_table = document.tables[0]
    assert title_table.rows[0].cells[0].paragraphs[0].runs[0].text == cca.name
    assert (
        "Compliance Criticality Assessment"
        in title_table.rows[1].cells[0].paragraphs[0].runs[0].text
    )

    solution_identification_table = document.tables[1]
    assert cca.solution_name in [
        solution_identification_table.rows[0].cells[1].text,
        None,
    ]
    assert cca.solution_version in [
        solution_identification_table.rows[1].cells[1].text,
        None,
    ]
    assert cca.vendor_name in [
        solution_identification_table.rows[2].cells[1].text,
        None,
    ]

    # Type of solution

    sub_table = solution_identification_table.rows[3].cells[1].tables[0]
    assert (
        checkbox[cca.solution_type == "application"] in sub_table.rows[0].cells[0].text
    )
    assert (
        checkbox[cca.solution_type == "middleware"] in sub_table.rows[0].cells[1].text
    )
    assert (
        checkbox[cca.solution_type == "infrastructure"]
        in sub_table.rows[1].cells[0].text
    )
    assert checkbox[cca.solution_type == "other"] in sub_table.rows[1].cells[1].text

    # hosting_type
    sub_table = solution_identification_table.rows[4].cells[1].tables[0]
    assert (
        checkbox[cca.solution_type == "third_party"] in sub_table.rows[0].cells[0].text
    )
    assert checkbox[cca.solution_type == "saas"] in sub_table.rows[0].cells[1].text
    assert (
        checkbox[cca.solution_type == "on_premises"] in sub_table.rows[1].cells[0].text
    )
    assert checkbox[cca.solution_type == "paas"] in sub_table.rows[1].cells[1].text
    assert checkbox[cca.solution_type == "website"] in sub_table.rows[2].cells[0].text
    assert checkbox[cca.solution_type == "iaas"] in sub_table.rows[2].cells[1].text

    # Server Host / Location
    assert cca.server_host in [
        solution_identification_table.rows[5].cells[1].text,
        None,
    ]

    # Solution Classification
    sub_table = solution_identification_table.rows[6].cells[1].tables[0]

    checkbox[cca.solution_classification == "custom"] in sub_table.rows[0].cells[0].text
    checkbox[cca.solution_classification == "configurable"] in sub_table.rows[1].cells[
        0
    ].text
    checkbox[cca.solution_classification == "non_configurable"] in sub_table.rows[
        2
    ].cells[0].text

    # GXP Impact
    table = document.tables[2]
    gxp_impact = cca.gxp_impact

    # row index
    assert table.rows[1].cells[0].text == "Good Manufacturing Practices (GMP) Impact"
    assert table.rows[2].cells[0].text == "1"
    assert table.rows[3].cells[0].text == "2"

    gxp_impact_keys = [
        (2, "is_solution_used_for_product_quality_control"),
        (3, "is_solution_part_of_batch_record"),
        (4, "is_impacted_by_gmp_global_regulations"),
        (6, "is_impacted_by_gcp_global_regulations"),
        (7, "is_solution_used_to_design_discover_products"),
        (9, "is_impacted_by_glp_global_regulations"),
        (10, "is_solution_used_to_collect_and_process_data"),
        (11, "is_solution_used_for_post_marketing_commitment"),
        (12, "is_solution_used_to_monitor_and_report_source_data"),
        (13, "is_solution_externally_facing_tool"),
        (15, "is_solution_used_to_make_quality_related_decisions"),
        (16, "is_solution_used_to_support_gxp_processes"),
    ]
    for index, key in gxp_impact_keys:
        assert checkbox_choice[gxp_impact.get(key)] in table.rows[index].cells[2].text

    # GXP ERES
    table = document.tables[3]
    gxp_eres = cca.gxp_eres
    gxp_eres_keys = [
        (
            1,
            "does_solution_create_records_in_electronic_form_required_by_gxp_regulation",
        ),
        (2, "does_solution_employ_electorinic_signatures"),
    ]

    for index, key in gxp_eres_keys:
        assert checkbox_choice[gxp_eres.get(key)] in table.rows[index].cells[2].text

    table = document.tables[4]
    sox_impact = cca.sox_impact

    sox_impact_keys = [
        (1, "is_solution_used_for_material_financial_data"),
        (2, "does_solution_provide_access_control_for_financial_systems"),
        (3, "does_system_feed_information_to_from_sox_system"),
    ]

    for index, key in sox_impact_keys:
        assert checkbox_choice[sox_impact.get(key)] in table.rows[index].cells[2].text

    privacy_impact = cca.privacy_impact
    table = document.tables[5]

    privacy_impact_keys = [
        (1, "will_personal_info_collected_be_from_individual_not_in_eu_country"),
        (2, "does_solution_collect_person_information"),
    ]

    for index, key in privacy_impact_keys:
        assert (
            checkbox_choice[privacy_impact.get(key)] in table.rows[index].cells[2].text
        )

    data_classification = cca.data_classification
    table = document.tables[6]

    data_classification_keys = [
        (1, "data_classification_secret"),
        (2, "data_classification_restricted"),
        (3, "data_classification_internal"),
        (4, "data_classification_public"),
    ]
    for index, key in data_classification_keys:
        assert checkbox[data_classification.get(key)] in table.rows[index].cells[2].text

    business_impact = cca.business_impact
    table = document.tables[7]

    business_impact_keys = [
        (2, "business_impact_high"),
        (3, "business_impact_medium"),
        (4, "business_impact_low"),
    ]
    for index, key in business_impact_keys:
        assert checkbox[business_impact.get(key)] in table.rows[index].cells[2].text

    business_impact = cca.business_impact
    table = document.tables[7]

    business_impact_keys = [
        (2, "business_impact_high"),
        (3, "business_impact_medium"),
        (4, "business_impact_low"),
    ]
    for index, key in business_impact_keys:
        assert checkbox[business_impact.get(key)] in table.rows[index].cells[2].text

    table = document.tables[10]
    rating = cca.rating
    rating_keys = [
        (1, "rating_significant"),
        (2, "rating_moderate"),
        (3, "rating_minimal"),
        (4, "rating_no_compliance_risk"),
    ]

    for index, key in rating_keys:
        assert checkbox[rating.get(key)] in table.rows[index].cells[2].text

    report = Report.objects.filter(cca=cca)
    assert report.count() == 1
    assert report[0].output_cca_doc is not None


def test_checkbox_helper():
    assert add_checkbox(True) == "☑"
    assert add_checkbox(False) == "☐"
